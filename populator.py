import pandas as pd
from docx import Document 
import sys
from datetime import datetime
import scraper 
import os

def populate(fileName):

    fileDir = os.path.join("Templates", fileName)
    fileDir += '.docx'
    document = Document(fileDir)

    date = datetime.today().strftime('%Y-%m-%d')
    pairs["[DATE]"] = date

    guarantor2 = True
    if pairs["[Guarantor2]"] == "nan" or pairs["[GuarantorAddress2]"] == "nan":
        guarantor2 = False



    if pairs["[GuarantorsAddress]"] == "nan":
        # Get corerspondence addresss
        corr_address = ""
        for line in personData["address"]:
            corr_address += personData["address"][line] + ", "
        corr_address = corr_address[:-2]

        pairs["[GuarantorsAddress]"] = corr_address

    if pairs["[Guarantors]"] == "nan":
        # get correctly ordered guarantor name
        oldName = personData["name"]
        split = oldName.split(",")
        name = split[1] + " " + split[0]
        pairs["[Guarantors]"] = name

    # get company name
    if pairs["[CompanyName]"] == "nan":
        pairs["[CompanyName]"] = scrapedData["company_name"]
    
    companyName = pairs["[CompanyName]"]


    if pairs["[PropertyAddressAndTitle]"] == "nan":
        # Get property address and title
        address = ""
        for line in scrapedData["registered_office_address"]:
            address += scrapedData["registered_office_address"][line] + ", "
        address = address[:-2]
        pairs["[PropertyAddressAndTitle]"] = address



    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        text = inline[i].text
                        extraText = "[" + text + "]"

                        if text in pairs.keys():
                            text=text.replace(text,pairs[text])
                            print("REPLACED with: " + text)
                            inline[i].text = text

                        if extraText in pairs.keys():
                            try:
                                extraText=extraText.replace(extraText,pairs[extraText])
                                print("REPLACED EXTRA with: " + extraText)
                                inline[i].text = extraText
                                inline[i+1].text = ""
                                inline[i-1].text = ""
                            except:
                                print("INDEX ERROR")
                    
    for p in document.paragraphs:
        inline = p.runs
        # print(p.text)
        for i in range(len(inline)):
            text = inline[i].text
            extraText = "[" + text + "]"


            if (text == "[Guarantor2]" or extraText == "[Guarantor2Address]") and guarantor2 == False:
                inline[i].text = ""
                inline[i-1].text = ""
                inline[i+1].text = ""
                continue
                
            if text in pairs.keys():
                text=text.replace(text,pairs[text])
                print("REPLACED with: " + text) 
                inline[i].text = text


            if extraText in pairs.keys():
                try:
                    extraText=extraText.replace(extraText,pairs[extraText])
                    print("REPLACED EXTRA with: " + extraText)
                    inline[i].text = extraText
                    inline[i+1].text = ""
                    inline[i-1].text = ""
                except:
                    print("INDEX ERROR")

    folderName = "Populated " + companyName

    if not os.path.exists(folderName):
        os.mkdir(folderName)

    outputdir = os.path.join(folderName , fileName)
    if dateOn:
        outputdir += ' Edited On ' + date + ".docx"
    else:
        outputdir += ' Edited' + ".docx"
    document.save(outputdir)
    print("--------------------------------------------------------------------------------------")





if __name__ == "__main__":
    arguments = sys.argv
    if len(arguments) <= 2:
        print("Not enough arguments, required format is [Company Number] [File to populate]")
        print("Note that the filename should match exactly without the extension (eg. .docx) included")
        sys.exit(1)
    
    dateOn = False
    if len(arguments) == 4:
        preferDate = str(arguments[3])
        if preferDate == 'date':
            dateOn = True   

    
    company_num = str(arguments[1])
    fileNameInput = arguments[2]

    # wSuDcPn_U0376euJz2zJmrZoYePPLXCjvM2OMuAJ    -- WHY CAN WE NOT USE THIS?
    grabber = scraper.CompaniesHouseService(key="uT7AYPcRf-CmKy5l-aCuALhKKn7vnR977Kr3NtQb")
    scrapedData = grabber.get_company_profile(company_num)
    if scrapedData == {}:
        print("-!-!-!-!-!-!-!-!-NO DATA WAS RETURNED-!-!-!-!-!-!-!-!-!")
    
    print("\n\n")
    print("----------------DETAILS OF COMPANY-------------------------------------")
    for data in scrapedData:
        print(str(data) + " : " + str(scrapedData[data]))

    print("\n\n----------------DETAILS OF DIRECTOR(S)-------------------------------------")
    personData = scrapedData["items"][0]
    for data in personData:
        print(str(data) + " : " + str(personData[data]))
            

    data_file = 'data.xlsx'
    data = pd.read_excel(data_file)
    pairs = data.to_dict()

    for pair in pairs:
        pairs[pair] = str(pairs[pair][0]) 

    print("\n\n----------------PREDEFINED DATA-------------------------------------")
    for pair in pairs:
        print(str(pair) + ":" + str(pairs[pair]))

    print("\n\n\n\n")
    print("----------------POPULATING FILE-------------------------------------")

    if fileNameInput == "all":
        populate("Personal Guarantee")
        populate("Assignment Letter")
        populate("Debenture")
        populate("Facility Agreement")
        populate("VAT Agency Agreement")
        populate("VAT Side Letter")
        populate("Deed Of Postponement")
    else:
        populate(fileNameInput)




